import re
import subprocess
import docx
str_= str(input("Enter the string: "))
list_word=str_.split()
l=[] 
doc=docx.Document()

# initialize new paragraph
p = doc.add_paragraph("")
#fo=open('file.docx','a')
for i in list_word:
    if i not in l:
        l.append(i)
for i in l:
        #fo.write(f"the word \"{i}\" appears {list_word.count(i)} times\n")
        ## add sentence to initialized paragraph
        p.add_run(str(f"the word \"{i}\" appears {list_word.count(i)} times\n"))
doc.save('file.docx')
open_word=subprocess.run("file.docx",shell= True)
after_str= re.sub("VB","BV",str_)
print(after_str)
